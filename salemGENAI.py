import streamlit as st
import streamlit.components.v1 as components
from google import genai
from google.genai import types
from PIL import Image
import io
import json
from fpdf import FPDF 
import pypdf
import docx
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# =============================================================================
# 1. Page Configuration & Title
# =============================================================================
st.set_page_config(page_title="SALEM GENAI", page_icon="🤖", layout="centered")

# --- CSS Styling for Layout Elements ---
st.markdown("""
    <style>
    .copy-container-wrapper {
        display: flex;
        justify-content: flex-end;
        margin-top: -8px;
        margin-bottom: 10px;
    }
    .stButton>button {
        text-overflow: ellipsis;
        white-space: nowrap;
        overflow: hidden;
    }
    </style>
""", unsafe_allow_html=True)

# --- Main Page Layout with Logo ---
col1, col2 = st.columns([1, 4]) 
with col1:
    try:
        st.image("salemlogo.png", width=80) 
    except Exception:
        st.write("🏫")
with col2:
    st.title("SALEM HILLS INT'L SCHOOL AI")

try:
    st.logo('salemlogo.png')
except Exception:
    pass

# =============================================================================
# Lesson Plan Helpers (PDF & Word Generators)
# =============================================================================
def clean_text(text: str) -> str:
    """Replaces common non-Latin-1 characters to prevent PDF generation errors."""
    replacements = {
        '\u201c': '"', '\u201d': '"', '\u2018': "'", '\u2019': "'",  
        '\u2013': '-', '\u2014': '-', '\u2022': '*',                  
    }
    for original, replacement in replacements.items():
        text = text.replace(original, replacement)
    return text.encode('latin-1', 'replace').decode('latin-1')

def generate_lesson_plan_pdf(lp: dict) -> bytes:
    """Generates a structured, clean, cell-populated tabular-like Lesson Plan PDF."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # Title / Header
    pdf.set_font("Helvetica", style="B", size=14)
    pdf.cell(0, 8, txt="SALEM HILLS INTERNATIONAL SCHOOL", ln=True, align="C")
    pdf.set_font("Helvetica", style="B", size=12)
    pdf.cell(0, 6, txt="LESSON PLAN - AUTUMN TERM (SESSION: 2026/27)", ln=True, align="C")
    pdf.ln(5)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)
    
    # Metadata Block
    pdf.set_font("Helvetica", size=10)
    pdf.write(5, f"Week: {clean_text(lp.get('Week',''))}   |   Date: {clean_text(lp.get('Date',''))}   |   Class: Year {clean_text(lp.get('Class',''))}\n")
    pdf.write(5, f"Period: {clean_text(lp.get('Period',''))}   |   Time: {clean_text(lp.get('Time',''))}   |   Duration: {clean_text(lp.get('Duration','40 Min'))}\n")
    pdf.write(5, f"Subject: {clean_text(lp.get('Subject','Mathematics'))}   |   Teacher: {clean_text(lp.get('Teacher','Ndubugo Njoagwuani'))}\n")
    pdf.ln(3)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)
    
    # Topic & Subtopic
    pdf.set_font("Helvetica", style="B", size=10)
    pdf.cell(30, 6, "Topic:", ln=False)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 6, clean_text(lp.get('Topic','')), ln=True)
    
    pdf.set_font("Helvetica", style="B", size=10)
    pdf.cell(30, 6, "Subtopic:", ln=False)
    pdf.set_font("Helvetica", size=10)
    pdf.cell(0, 6, clean_text(lp.get('Subtopic','')), ln=True)
    pdf.ln(2)
    
    # Objectives & Knowledge Blocks
    sections = [
        ("Behavioural Objectives", "At the end of the lesson, the students should be able to:\n" + lp.get('Behavioural Objectives','')),
        ("Previous Knowledge", "The students are already familiar with:\n" + lp.get('Previous Knowledge','')),
        ("Instructional Resources", lp.get('Instructional Resources','')),
        ("Key Words / Start", lp.get('Key Words',''))
    ]
    
    for title, content in sections:
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(0, 6, f"{title}:", ln=True)
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5, clean_text(content))
        pdf.ln(2)
        
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)
    
    # Presentation Table Block
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(0, 6, "Presentation Steps:", ln=True)
    pdf.ln(2)
    
    steps = ["Step 1", "Step 2", "Step 3"]
    for step in steps:
        step_data = lp.get("Presentation", {}).get(step, {})
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(0, 6, f"--- {step} ---", ln=True)
        
        pdf.set_font("Helvetica", style="I", size=10)
        pdf.cell(40, 5, "Teacher's Activity:", ln=False)
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5, clean_text(step_data.get("Teacher Activity", "")))
        
        pdf.set_font("Helvetica", style="I", size=10)
        pdf.cell(40, 5, "Students' Activity:", ln=False)
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5, clean_text(step_data.get("Students Activity", "")))
        pdf.ln(2)
        
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)
    
    # Closure Blocks
    closures = [
        ("Summary", lp.get('Summary','')),
        ("Evaluation", lp.get('Evaluation','')),
        ("Conclusion", lp.get('Conclusion','')),
        ("Homework", lp.get('Homework',''))
    ]
    for title, content in closures:
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(30, 6, f"{title}:", ln=False)
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5, clean_text(content))
        pdf.ln(1)
        
    pdf.ln(5)
    pdf.set_font("Helvetica", style="B", size=10)
    pdf.cell(0, 6, "Sectional Head’s Comment: ________________________  Signature: _________  Date: _________", ln=True)
    
    return bytes(pdf.output())

def set_cell_background(cell, fill_hex):
    """Applies clean background color formatting to Word Document cells."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_lesson_plan_docx(lp: dict) -> bytes:
    """Generates an accurate duplication of the template as a fully populated Word table."""
    doc = docx.Document()
    
    # Set standard Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    title = doc.add_paragraph()
    run = title.add_run("SALEM HILLS INTERNATIONAL SCHOOL\nLESSON PLAN")
    run.bold = True
    font_t = run.font
    font_t.size = Pt(14)
    title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    meta_p = doc.add_paragraph()
    meta_p.add_run("Term: AUTUMN   |   Session: 2026/27\n").bold = True
    meta_p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    
    # Create the Master Table Frame Blueprint
    table = doc.add_table(rows=14, cols=2)
    table.style = 'Table Grid'
    
    # Top Meta Grid Rows
    row0 = table.rows[0].cells
    row0[0].text = f"Week: {lp.get('Week','')}\nDate: {lp.get('Date','')}"
    row0[1].text = f"Class: Year {lp.get('Class','')}\nPeriod: {lp.get('Period','')}"
    
    row1 = table.rows[1].cells
    row1[0].text = f"Time: {lp.get('Time','')}\nDuration: {lp.get('Duration','40 Min')}"
    row1[1].text = f"Subject: {lp.get('Subject','Mathematics')}\nTeacher: {lp.get('Teacher','Ndubugo Njoagwuani')}"
    
    # Core Lesson Content Rows
    table.rows[2].cells[0].text = f"Topic:\n{lp.get('Topic','')}"
    table.rows[2].cells[1].text = f"Subtopic:\n{lp.get('Subtopic','')}"
    
    table.rows[3].cells[0].text = f"Behavioural Objectives:\nAt the end of the lesson, the students should be able to:\n{lp.get('Behavioural Objectives','')}"
    table.rows[3].cells[1].text = f"Previous Knowledge:\nThe students are already familiar with...\n{lp.get('Previous Knowledge','')}"
    
    table.rows[4].cells[0].text = f"Instructional Resources:\n{lp.get('Instructional Resources','')}"
    table.rows[4].cells[1].text = f"Key words / Start:\n{lp.get('Key Words','')}"
    
    # Presentation Rows Header
    cell_pres_title = table.cell(5, 0).merge(table.cell(5, 1))
    cell_pres_title.text = "Presentation Details"
    set_cell_background(cell_pres_title, "EFEFEF")
    cell_pres_title.paragraphs[0].runs[0].font.bold = True
    
    steps = ["Step 1", "Step 2", "Step 3"]
    row_idx = 6
    for step in steps:
        step_data = lp.get("Presentation", {}).get(step, {})
        table.cell(row_idx, 0).text = f"{step} - Teacher’s Activity:\n{step_data.get('Teacher Activity','')}"
        table.cell(row_idx, 1).text = f"{step} - Students’ Activity:\n{step_data.get('Students Activity','')}"
        row_idx += 1
        
    # Closing Blocks Rows
    table.cell(9, 0).text = f"Summary:\n{lp.get('Summary','')}"
    table.cell(9, 1).text = f"Evaluation:\n{lp.get('Evaluation','')}"
    
    table.cell(10, 0).text = f"Conclusion:\n{lp.get('Conclusion','')}"
    table.cell(10, 1).text = f"Homework:\n{lp.get('Homework','')}"
    
    # Approvals Footer Rows
    cell_foot = table.cell(11, 0).merge(table.cell(11, 1))
    cell_foot.text = "Sectional Head’s Comment:"
    
    table.cell(12, 0).text = "Signature:"
    table.cell(12, 1).text = "Date:"
    
    # Adjust widths for standard layout look
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(3.25)
            
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

# =============================================================================
# Core PDF Transcription & Global Parsing Functions
# =============================================================================
def generate_chat_pdf(messages) -> bytes:
    """Generates a regular transcript conversation log PDF file."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Helvetica", style="B", size=14)
    pdf.cell(0, 10, txt="SALEM HILLS INT'L SCHOOL CHAT LOG", ln=True, align="C")
    pdf.ln(5)
    
    for msg in messages:
        role = "Student / User" if msg["role"] == "user" else "AI Assistant"
        pdf.set_font("Helvetica", style="B", size=10)
        pdf.cell(0, 6, txt=f"{role}:", ln=True)
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(0, 5, txt=clean_text(msg["content"]))
        pdf.ln(4)
    return bytes(pdf.output())

def extract_text_from_file(uploaded_file) -> str:
    """Extracts text strings from uploaded txt, pdf, and docx variants."""
    if uploaded_file.name.endswith('.txt'):
        return uploaded_file.read().decode("utf-8")
    elif uploaded_file.name.endswith('.pdf'):
        try:
            pdf_reader = pypdf.PdfReader(uploaded_file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        except Exception as e:
            return f"[Error: {e}]"
    elif uploaded_file.name.endswith('.docx'):
        try:
            doc = docx.Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            return f"[Error: {e}]"
    return ""

def render_copy_button(text_to_copy: str, element_key: str):
    """Renders JavaScript powered clipboard functionality inside chat messages."""
    safe_text = text_to_copy.replace('\\', '\\\\').replace('`', '\\`').replace('$', '\\$')
    html_code = f"""
    <div style="display: flex; justify-content: flex-end; font-family: sans-serif;">
        <button onclick="copyToClipboard()" id="btn_{element_key}" style="
            background: none; border: none; color: #6c757d; cursor: pointer;
            display: flex; align-items: center; gap: 4px; font-size: 12px;
            padding: 4px 8px; border-radius: 4px; transition: all 0.2s;
        " onmouseover="this.style.color='#000'; this.style.background='#f0f2f6'" 
           onmouseout="this.style.color='#6c757d'; this.style.background='none'">
            <svg xmlns="http://www.w3.org/2000/svg" width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><rect x="9" y="9" width="13" height="13" rx="2" ry="2"></rect><path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"></path></svg>
            <span id="label_{element_key}">Copy Prompt</span>
        </button>
    </div>
    <script>
    function copyToClipboard() {{
        const text = `{safe_text}`;
        navigator.clipboard.writeText(text).then(() => {{
            const label = document.getElementById('label_{element_key}');
            const btn = document.getElementById('btn_{element_key}');
            label.innerText = 'Copied!'; btn.style.color = '#28a745';
            setTimeout(() => {{ label.innerText = 'Copy Prompt'; btn.style.color = '#6c757d'; }}, 2000);
        }});
    }}
    </script>
    """
    st.markdown('<div class="copy-container-wrapper">', unsafe_allow_html=True)
    components.html(html_code, height=30)
    st.markdown('</div>', unsafe_allow_html=True)

# =============================================================================
# 3. Memory & Client Configurations
# =============================================================================
if "messages" not in st.session_state:
    st.session_state.messages = []
if "input_value" not in st.session_state:
    st.session_state.input_value = ""
if "current_lesson_plan" not in st.session_state:
    st.session_state.current_lesson_plan = None

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
else:
    api_key = None

client = None
if api_key:
    try:
        client = genai.Client(api_key=api_key, http_options={"api_version": "v1beta"})
    except Exception as e:
        st.error(f"Failed to initialize AI client: {e}")

# =============================================================================
# 4. Sidebar Panels
# =============================================================================
with st.sidebar:
    st.header("Configuration")
    app_mode = st.radio("Select App Mode:", ("💬 Text Chat", "📋 Lesson Plan Generator"))
    st.divider()

    if app_mode == "💬 Text Chat":
        st.subheader("Text Mode Settings")
        system_instruction = st.text_area(
            "System Instructions / AI Persona:",
            value="You are a helpful, expert educational assistant. Break down complex topics simply."
        )
        st.divider()
        st.subheader("Query History")
        user_prompts = []
        for msg in st.session_state.messages:
            if msg["role"] == "user":
                clean_string = msg["content"].split("\n\n")[-1] if "📄 *Attached file:" in msg["content"] else msg["content"]
                if clean_string not in user_prompts:
                    user_prompts.append(clean_string)
        
        if user_prompts:
            for i, prompt in enumerate(user_prompts):
                short_display = prompt[:30] + "..." if len(prompt) > 30 else prompt
                if st.button(f"💬 {short_display}", key=f"hist_btn_{i}", use_container_width=True):
                    st.session_state.input_value = prompt
                    st.rerun()
        else:
            st.info("No queries sent yet.")
            
        st.divider()
        st.subheader("Export Conversation")
        has_messages = len(st.session_state.messages) > 0
        if st.download_button(label="📥 Download Chat as PDF", data=generate_chat_pdf(st.session_state.messages) if has_messages else b"", file_name="chat_transcript.pdf", mime="application/pdf", disabled=not has_messages):
            st.success("Log Transcribed!")
            
    elif app_mode == "📋 Lesson Plan Generator":
        st.subheader("Export Documents")
        if st.session_state.current_lesson_plan:
            lp_data = st.session_state.current_lesson_plan
            
            pdf_data = generate_lesson_plan_pdf(lp_data)
            st.download_button(
                label="📥 Download Plan as PDF (Portrait)",
                data=pdf_data,
                file_name=f"Lesson_Plan_{lp_data.get('Topic','Plan')}.pdf",
                mime="application/pdf",
                use_container_width=True
            )
            
            docx_data = generate_lesson_plan_docx(lp_data)
            st.download_button(
                label="📥 Download Plan as Word (.docx)",
                data=docx_data,
                file_name=f"Lesson_Plan_{lp_data.get('Topic','Plan')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.info("Generate a lesson plan first to enable downloads.")

    st.divider()
    if st.button("Clear App History / Cache"):
        st.session_state.messages = []
        st.session_state.input_value = ""
        st.session_state.current_lesson_plan = None
        st.rerun()

# =============================================================================
# 5. Application Core Context Router
# =============================================================================
if app_mode == "💬 Text Chat":
    for idx, message in enumerate(st.session_state.messages):
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
        if message["role"] == "user":
            raw_text = message["content"].split("\n\n")[-1] if "📄 *Attached file:" in message["content"] else message["content"]
            render_copy_button(raw_text, f"hist_{idx}")

    uploaded_doc = st.file_uploader("Attach a document (.txt, .pdf, or .docx)", type=["txt", "pdf", "docx"], label_visibility="collapsed")
    user_input = st.chat_input("Ask SALEM anything...")

    if st.session_state.input_value and not user_input:
        user_input = st.session_state.input_value
        st.session_state.input_value = ""

    if user_input:
        if not api_key:
            st.warning("Please provide a valid Gemini API Key in the sidebar secrets configuration to start.")
            st.stop()
        
        doc_text = ""
        if uploaded_doc is not None:
            with st.spinner("Extracting content..."):
                doc_text = extract_text_from_file(uploaded_doc)
        
        full_prompt = user_input
        display_prompt = user_input
        if doc_text:
            full_prompt = f"User Message: {user_input}\n\nAttached Document Content:\n{doc_text}"
            display_prompt = f"📄 *Attached file: {uploaded_doc.name}*\n\n{user_input}"

        with st.chat_message("user"):
            st.markdown(display_prompt)
        render_copy_button(user_input, "live_new")
        st.session_state.messages.append({"role": "user", "content": display_prompt})

        with st.chat_message("assistant"):
            message_placeholder = st.empty()
            try:
                with st.spinner("Processing..."):
                    chat = client.chats.create(
                        model="gemini-2.5-flash",
                        config=types.GenerateContentConfig(system_instruction=system_instruction, temperature=0.7)
                    )
                    response = chat.send_message(full_prompt)
                message_placeholder.markdown(response.text)
                st.session_state.messages.append({"role": "assistant", "content": response.text})
                st.rerun()
            except Exception as e:
                st.error(f"Error: {e}")

# =============================================================================
# 6. Mode B: Lesson Plan Generator Screen
# =============================================================================
elif app_mode == "📋 Lesson Plan Generator":
    st.subheader("Generate Structural Template Lesson Plans")
    st.caption("Populates specific matrix layouts for Salem Hills International School matching your structural curriculum grids.")
    
    col_l, col_r = st.columns(2)
    with col_l:
        target_topic = st.text_input("Enter Target Topic:", placeholder="e.g., Quadratic Equations")
        target_subtopic = st.text_input("Enter Subtopic:", placeholder="e.g., Factorization Method")
    with col_r:
        target_class = st.text_input("Year / Class:", value="10")
        target_week = st.text_input("Term Week #:", value="3")

    if st.button("Generate Template Lesson Plan", type="primary"):
        if not target_topic:
            st.error("Please provide at least a primary target topic description.")
        else:
            if not api_key:
                st.warning("Please configure your Gemini API Key first.")
                st.stop()
                
            lesson_plan_prompt = f"""
            You are an expert curriculum supervisor tasked with populating the Salem Hills International School Lesson Plan Template.
            Create a comprehensive lesson plan based on:
            Topic: {target_topic}
            Subtopic: {target_subtopic}
            Class: Year {target_class}
            Week: {target_week}
            Teacher: Ndubugo Njoagwuani
            Subject: Mathematics
            Duration: 40 Min
            
            You MUST respond ONLY with a raw JSON object that maps perfectly to these keys. Do not wrap it in markdown codeblocks. 
            Ensure every column field is detailed, practical, and highly academic.
            
            Expected JSON Blueprint Format:
            {{
                "Week": "{target_week}",
                "Date": "Auto",
                "Class": "{target_class}",
                "Period": "1 & 2",
                "Time": "08:00 AM",
                "Duration": "40 Min",
                "Subject": "Mathematics",
                "Teacher": "Ndubugo Njoagwuani",
                "Topic": "{target_topic}",
                "Subtopic": "{target_subtopic}",
                "Behavioural Objectives": "[Provide clear bullet items outlining what the students will be able to perform]",
                "Previous Knowledge": "[Detail explicitly what foundational material the students are already familiar with]",
                "Instructional Resources": "[List textbook details, grid tools, chalkboard requirements, or smart board setups]",
                "Key Words": "[List core vocabulary key terms for this session]",
                "Presentation": {{
                    "Step 1": {{
                        "Teacher Activity": "[Detailed description of step 1 introduction activities]",
                        "Students Activity": "[What students are actively doing during step 1]"
                    }},
                    "Step 2": {{
                        "Teacher Activity": "[Core presentation mechanics of step 2 content delivery]",
                        "Students Activity": "[What students are doing during core delivery]"
                    }},
                    "Step 3": {{
                        "Teacher Activity": "[Classwork exercise monitoring or problem evaluation guidance details]",
                        "Students Activity": "[Student independent practice activities]"
                    }}
                }},
                "Summary": "[Brief high-level wrap-up sentence synthesis]",
                "Evaluation": "[List specific conceptual checking validation assessment questions]",
                "Conclusion": "[Final closing action steps summary note]",
                "Homework": "[Provide structured practice assignment problems]"
            }}
            """
            
            with st.spinner("AI is structuring your academic template..."):
                try:
                    response = client.models.generate_content(
                        model="gemini-2.5-flash",
                        contents=lesson_plan_prompt,
                        config=types.GenerateContentConfig(
                            response_mime_type="application/json",
                            temperature=0.3
                        )
                    )
                    
                    # Store the parsed dictionary directly into the state
                    st.session_state.current_lesson_plan = json.loads(response.text)
                    st.success("Lesson Plan populated successfully! Use the sidebar panel on the left to export files.")
                except Exception as e:
                    st.error(f"Failed to generate lesson structure data: {e}")
                    
    # Render the structured live preview of the plan if it exists in state
    if st.session_state.current_lesson_plan:
        lp = st.session_state.current_lesson_plan
        st.markdown("---")
        st.subheader("📋 Document Grid Matrix Live Preview")
        
        # Display metadata fields
        st.info(f"**{lp.get('Subject','Mathematics')} Plan** | Week {lp.get('Week','')} | Year Class: {lp.get('Class','')}")
        
        # Display contents in layout blocks matching the paper cells
        st.markdown(f"### **Topic:** {lp.get('Topic','')}")
        st.markdown(f"**Subtopic:** {lp.get('Subtopic','')}")
        
        col_obj, col_prev = st.columns(2)
        with col_obj:
            st.markdown("**Behavioural Objectives:**")
            st.write(lp.get('Behavioural Objectives',''))
        with col_prev:
            st.markdown("**Previous Knowledge:**")
            st.write(lp.get('Previous Knowledge',''))
            
        st.markdown("**Instructional Resources:**")
        st.caption(lp.get('Instructional Resources',''))
        
        st.markdown("#### **Presentation Delivery Matrix**")
        for step in ["Step 1", "Step 2", "Step 3"]:
            with st.expander(f"➔ {step} Details", expanded=True):
                st.write(f"**Teacher's Activity:** {lp.get('Presentation', {}).get(step, {}).get('Teacher Activity','')}")
                st.write(f"**Students' Activity:** {lp.get('Presentation', {}).get(step, {}).get('Students Activity','')}")
                
        st.markdown("#### **Closure Matrix Summary**")
        st.write(f"**Evaluation questions:** {lp.get('Evaluation','')}")
        st.write(f"**Homework:** {lp.get('Homework','')}")
